"""
Document processor for extracting text from various file formats.
Supports: PDF, Word (DOCX), CSV, Excel (XLSX), and plain text.
"""

import io
import logging
from typing import Optional
from pathlib import Path

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pandas as pd
except ImportError:
    pd = None

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process various document formats and extract text."""
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file."""
        if PdfReader is None:
            raise ImportError("PyPDF2 is required for PDF processing. Install with: pip install pypdf2")
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"--- Page {page_num} ---\n{text}")
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise ValueError(f"Failed to process PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from Word DOCX file."""
        if Document is None:
            raise ImportError("python-docx is required for Word processing. Install with: pip install python-docx")
        
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise ValueError(f"Failed to process Word document: {str(e)}")
    
    @staticmethod
    def extract_text_from_csv(file_content: bytes) -> str:
        """Extract text from CSV file."""
        if pd is None:
            raise ImportError("pandas is required for CSV processing. Install with: pip install pandas")
        
        try:
            csv_file = io.BytesIO(file_content)
            df = pd.read_csv(csv_file)
            
            # Convert DataFrame to readable text format
            text_parts = [f"CSV Data ({len(df)} rows, {len(df.columns)} columns)\n"]
            text_parts.append("Columns: " + ", ".join(df.columns))
            text_parts.append("\n--- Data ---")
            
            # Convert to string representation
            text_parts.append(df.to_string(index=False))
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from CSV: {e}")
            raise ValueError(f"Failed to process CSV: {str(e)}")
    
    @staticmethod
    def extract_text_from_excel(file_content: bytes) -> str:
        """Extract text from Excel XLSX file."""
        if pd is None:
            raise ImportError("pandas is required for Excel processing. Install with: pip install pandas openpyxl")
        
        try:
            excel_file = io.BytesIO(file_content)
            # Read all sheets
            excel_data = pd.read_excel(excel_file, sheet_name=None)
            
            text_parts = []
            for sheet_name, df in excel_data.items():
                text_parts.append(f"\n=== Sheet: {sheet_name} ===")
                text_parts.append(f"({len(df)} rows, {len(df.columns)} columns)")
                text_parts.append("Columns: " + ", ".join(df.columns))
                text_parts.append("\n--- Data ---")
                text_parts.append(df.to_string(index=False))
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from Excel: {e}")
            raise ValueError(f"Failed to process Excel file: {str(e)}")
    
    @staticmethod
    def process_file(filename: str, file_content: bytes) -> str:
        """
        Process a file and extract text based on file extension.
        
        Args:
            filename: Name of the file
            file_content: Binary content of the file
            
        Returns:
            Extracted text content
        """
        file_ext = Path(filename).suffix.lower()
        
        if file_ext == '.pdf':
            return DocumentProcessor.extract_text_from_pdf(file_content)
        elif file_ext in ['.docx', '.doc']:
            return DocumentProcessor.extract_text_from_docx(file_content)
        elif file_ext == '.csv':
            return DocumentProcessor.extract_text_from_csv(file_content)
        elif file_ext in ['.xlsx', '.xls']:
            return DocumentProcessor.extract_text_from_excel(file_content)
        elif file_ext == '.txt':
            return file_content.decode('utf-8')
        else:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported formats: PDF, DOCX, CSV, XLSX, TXT")

